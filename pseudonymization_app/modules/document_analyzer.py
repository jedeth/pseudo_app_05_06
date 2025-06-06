#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Analyseur de documents Word pour l'application de pseudonymisation
================================================================

Ce module permet d'analyser, extraire et traiter des documents Word (.docx)
pour la pseudonymisation. Il maintient la structure et le formatage du document
tout en permettant l'extraction et la modification du contenu textuel.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Tuple, Any, Optional
from datetime import datetime
import zipfile
import xml.etree.ElementTree as ET

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ Module python-docx non installé. Installez-le avec: pip install python-docx")

class DocumentAnalyzer:
    """
    Analyseur de documents Word pour la pseudonymisation
    
    Cette classe permet de :
    - Lire et analyser des documents .docx
    - Extraire le texte en préservant la structure
    - Identifier les entités à pseudonymiser
    - Modifier le document en conservant le formatage
    - Sauvegarder les documents pseudonymisés
    """
    
    def __init__(self):
        """
        Initialise l'analyseur de documents
        """
        if not DOCX_AVAILABLE:
            raise ImportError("Le module python-docx est requis. Installez-le avec: pip install python-docx")
        
        self.document = None
        self.original_path = ""
        self.extracted_text = ""
        self.document_structure = []
        self.entities_found = {}
        
    def load_document(self, filepath: str) -> bool:
        """
        Charge un document Word depuis un fichier
        
        Args:
            filepath (str): Chemin vers le fichier .docx
            
        Returns:
            bool: True si le chargement réussit, False sinon
        """
        try:
            # Vérifie que le fichier existe
            if not Path(filepath).exists():
                raise FileNotFoundError(f"Le fichier {filepath} n'existe pas")
            
            # Vérifie l'extension
            if not filepath.lower().endswith('.docx'):
                raise ValueError("Seuls les fichiers .docx sont supportés")
            
            # Charge le document
            self.document = Document(filepath)
            self.original_path = filepath
            
            print(f"✅ Document chargé: {Path(filepath).name}")
            return True
            
        except Exception as e:
            print(f"❌ Erreur lors du chargement du document: {e}")
            return False
    
    def extract_text_with_structure(self) -> Dict[str, Any]:
        """
        Extrait le texte du document en préservant la structure
        
        Returns:
            Dict: Informations sur le document et son contenu
        """
        if not self.document:
            raise ValueError("Aucun document chargé")
        
        document_info = {
            'filename': Path(self.original_path).name,
            'total_paragraphs': 0,
            'total_tables': 0,
            'total_words': 0,
            'paragraphs': [],
            'tables': [],
            'headers_footers': [],
            'full_text': ""
        }
        
        all_text_parts = []
        
        # Extraction des paragraphes
        print("📖 Extraction des paragraphes...")
        for i, paragraph in enumerate(self.document.paragraphs):
            if paragraph.text.strip():  # Ignore les paragraphes vides
                paragraph_info = {
                    'index': i,
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else 'Normal',
                    'alignment': str(paragraph.alignment) if paragraph.alignment else 'Left',
                    'word_count': len(paragraph.text.split())
                }
                
                document_info['paragraphs'].append(paragraph_info)
                all_text_parts.append(paragraph.text)
        
        document_info['total_paragraphs'] = len(document_info['paragraphs'])
        
        # Extraction des tableaux
        print("📋 Extraction des tableaux...")
        for i, table in enumerate(self.document.tables):
            table_info = {
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'cells': []
            }
            
            for row_idx, row in enumerate(table.rows):
                row_cells = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append({
                            'row': row_idx,
                            'col': col_idx,
                            'text': cell_text
                        })
                        all_text_parts.append(cell_text)
                
                if row_cells:
                    table_info['cells'].extend(row_cells)
            
            if table_info['cells']:
                document_info['tables'].append(table_info)
        
        document_info['total_tables'] = len(document_info['tables'])
        
        # Extraction des en-têtes et pieds de page
        print("📄 Extraction des en-têtes et pieds de page...")
        for section in self.document.sections:
            # En-têtes
            if section.header.paragraphs:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        document_info['headers_footers'].append({
                            'type': 'header',
                            'text': paragraph.text
                        })
                        all_text_parts.append(paragraph.text)
            
            # Pieds de page
            if section.footer.paragraphs:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        document_info['headers_footers'].append({
                            'type': 'footer',
                            'text': paragraph.text
                        })
                        all_text_parts.append(paragraph.text)
        
        # Texte complet
        document_info['full_text'] = '\n'.join(all_text_parts)
        document_info['total_words'] = len(document_info['full_text'].split())
        
        # Sauvegarde pour usage ultérieur
        self.extracted_text = document_info['full_text']
        self.document_structure = document_info
        
        print(f"✅ Extraction terminée: {document_info['total_paragraphs']} paragraphes, "
              f"{document_info['total_tables']} tableaux, {document_info['total_words']} mots")
        
        return document_info
    
    def analyze_entities_in_document(self, nlp_model, entity_types: List[str] = None) -> Dict[str, Any]:
        """
        Analyse les entités NER dans le document
        
        Args:
            nlp_model: Modèle SpaCy chargé pour l'analyse NER
            entity_types: Types d'entités à rechercher (None = tous)
            
        Returns:
            Dict: Résultats de l'analyse des entités
        """
        if not self.extracted_text:
            raise ValueError("Aucun texte extrait. Appelez extract_text_with_structure() d'abord")
        
        print("🔍 Analyse des entités dans le document...")
        
        # Analyse avec SpaCy
        doc = nlp_model(self.extracted_text)
        
        # Collecte des entités
        entities_analysis = {
            'total_entities': 0,
            'entities_by_type': {},
            'entities_by_location': {},
            'unique_entities': set(),
            'detailed_entities': []
        }
        
        for ent in doc.ents:
            # Filtre par type si spécifié
            if entity_types and ent.label_ not in entity_types:
                continue
            
            entity_info = {
                'text': ent.text,
                'label': ent.label_,
                'start': ent.start_char,
                'end': ent.end_char,
                'confidence': getattr(ent, 'score', 1.0)  # Certains modèles n'ont pas de score
            }
            
            # Ajoute aux statistiques
            entities_analysis['detailed_entities'].append(entity_info)
            entities_analysis['unique_entities'].add(ent.text)
            
            # Comptage par type
            if ent.label_ not in entities_analysis['entities_by_type']:
                entities_analysis['entities_by_type'][ent.label_] = []
            entities_analysis['entities_by_type'][ent.label_].append(ent.text)
            
            # Localisation dans le document
            location = self._find_entity_location(ent.text, ent.start_char)
            if location:
                if location['type'] not in entities_analysis['entities_by_location']:
                    entities_analysis['entities_by_location'][location['type']] = []
                entities_analysis['entities_by_location'][location['type']].append({
                    'entity': ent.text,
                    'label': ent.label_,
                    'location_details': location
                })
        
        entities_analysis['total_entities'] = len(entities_analysis['detailed_entities'])
        entities_analysis['unique_entities'] = list(entities_analysis['unique_entities'])
        
        # Sauvegarde pour usage ultérieur
        self.entities_found = entities_analysis
        
        print(f"✅ Analyse terminée: {entities_analysis['total_entities']} entités trouvées, "
              f"{len(entities_analysis['unique_entities'])} uniques")
        
        return entities_analysis
    
    def _find_entity_location(self, entity_text: str, char_position: int) -> Optional[Dict[str, Any]]:
        """
        Trouve la localisation d'une entité dans la structure du document
        
        Args:
            entity_text: Texte de l'entité
            char_position: Position du caractère dans le texte complet
            
        Returns:
            Dict: Informations sur la localisation ou None
        """
        if not self.document_structure:
            return None
        
        current_position = 0
        
        # Recherche dans les paragraphes
        for paragraph in self.document_structure['paragraphs']:
            paragraph_end = current_position + len(paragraph['text']) + 1  # +1 pour le \n
            
            if current_position <= char_position < paragraph_end:
                return {
                    'type': 'paragraph',
                    'index': paragraph['index'],
                    'style': paragraph['style'],
                    'relative_position': char_position - current_position
                }
            
            current_position = paragraph_end
        
        # Recherche dans les tableaux
        for table in self.document_structure['tables']:
            for cell in table['cells']:
                cell_end = current_position + len(cell['text']) + 1
                
                if current_position <= char_position < cell_end:
                    return {
                        'type': 'table_cell',
                        'table_index': table['index'],
                        'row': cell['row'],
                        'col': cell['col'],
                        'relative_position': char_position - current_position
                    }
                
                current_position = cell_end
        
        return None
    
    def pseudonymize_document(self, pseudonymization_map: Dict[str, str], 
                            highlight_changes: bool = True) -> Document:
        """
        Crée une version pseudonymisée du document
        
        Args:
            pseudonymization_map: Dictionnaire {texte_original: texte_pseudonymisé}
            highlight_changes: Surligner les modifications en jaune
            
        Returns:
            Document: Nouveau document pseudonymisé
        """
        if not self.document:
            raise ValueError("Aucun document chargé")
        
        print("🔄 Pseudonymisation du document en cours...")
        
        # Crée une copie du document
        pseudonymized_doc = Document()
        
        # Copie les styles du document original (simplifié)
        try:
            pseudonymized_doc.styles = self.document.styles
        except:
            pass  # Ignore les erreurs de copie de styles
        
        # Traite chaque paragraphe
        for paragraph in self.document.paragraphs:
            new_paragraph = pseudonymized_doc.add_paragraph()
            
            # Applique le pseudonymisation au texte
            pseudonymized_text = self._apply_pseudonymization(paragraph.text, pseudonymization_map)
            
            # Ajoute le texte au nouveau paragraphe
            if pseudonymized_text != paragraph.text and highlight_changes:
                # Surligne les modifications
                run = new_paragraph.add_run(pseudonymized_text)
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except:
                    pass  # Ignore si le surlignage échoue
            else:
                new_paragraph.add_run(pseudonymized_text)
            
            # Copie le style si possible
            try:
                new_paragraph.style = paragraph.style
            except:
                pass
        
        # Traite les tableaux
        for table in self.document.tables:
            new_table = pseudonymized_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    original_text = cell.text
                    pseudonymized_text = self._apply_pseudonymization(original_text, pseudonymization_map)
                    
                    new_cell = new_table.cell(row_idx, col_idx)
                    new_cell.text = pseudonymized_text
                    
                    # Surligne si modifié
                    if pseudonymized_text != original_text and highlight_changes:
                        try:
                            for paragraph in new_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        except:
                            pass
        
        print("✅ Pseudonymisation du document terminée")
        return pseudonymized_doc
    
    def _apply_pseudonymization(self, text: str, pseudonymization_map: Dict[str, str]) -> str:
        """
        Applique la pseudonymisation à un texte donné
        
        Args:
            text: Texte original
            pseudonymization_map: Dictionnaire de correspondance
            
        Returns:
            str: Texte pseudonymisé
        """
        pseudonymized_text = text
        
        # Applique chaque remplacement
        for original, pseudonym in pseudonymization_map.items():
            # Utilise une recherche sensible à la casse avec préservation du contexte
            pattern = re.compile(re.escape(original), re.IGNORECASE)
            pseudonymized_text = pattern.sub(pseudonym, pseudonymized_text)
        
        return pseudonymized_text
    
    def save_pseudonymized_document(self, pseudonymized_doc: Document, 
                                  output_path: str = None) -> str:
        """
        Sauvegarde le document pseudonymisé
        
        Args:
            pseudonymized_doc: Document pseudonymisé
            output_path: Chemin de sortie (généré automatiquement si None)
            
        Returns:
            str: Chemin du fichier sauvegardé
        """
        if output_path is None:
            # Génère un nom de fichier automatique
            original_name = Path(self.original_path).stem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"data/{original_name}_pseudonymized_{timestamp}.docx"
        
        # Crée le dossier de sortie si nécessaire
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        try:
            pseudonymized_doc.save(output_path)
            print(f"💾 Document pseudonymisé sauvegardé: {output_path}")
            return output_path
        except Exception as e:
            raise Exception(f"Erreur lors de la sauvegarde: {e}")
    
    def export_analysis_report(self, output_path: str = None) -> str:
        """
        Exporte un rapport d'analyse au format texte
        
        Args:
            output_path: Chemin de sortie du rapport
            
        Returns:
            str: Chemin du fichier de rapport
        """
        if not self.document_structure or not self.entities_found:
            raise ValueError("Aucune analyse disponible")
        
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"data/analysis_report_{timestamp}.txt"
        
        # Crée le dossier de sortie si nécessaire
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        report_content = []
        report_content.append("RAPPORT D'ANALYSE DE DOCUMENT")
        report_content.append("=" * 50)
        report_content.append(f"Document: {self.document_structure['filename']}")
        report_content.append(f"Date d'analyse: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report_content.append("")
        
        # Statistiques générales
        report_content.append("STATISTIQUES GÉNÉRALES")
        report_content.append("-" * 30)
        report_content.append(f"Nombre de paragraphes: {self.document_structure['total_paragraphs']}")
        report_content.append(f"Nombre de tableaux: {self.document_structure['total_tables']}")
        report_content.append(f"Nombre total de mots: {self.document_structure['total_words']}")
        report_content.append("")
        
        # Entités trouvées
        report_content.append("ENTITÉS IDENTIFIÉES")
        report_content.append("-" * 30)
        report_content.append(f"Total d'entités: {self.entities_found['total_entities']}")
        report_content.append(f"Entités uniques: {len(self.entities_found['unique_entities'])}")
        report_content.append("")
        
        # Détail par type d'entité
        for entity_type, entities in self.entities_found['entities_by_type'].items():
            report_content.append(f"{entity_type}: {len(entities)} occurrences")
            unique_entities = list(set(entities))
            for entity in unique_entities[:10]:  # Limite à 10 exemples
                count = entities.count(entity)
                report_content.append(f"  - {entity} ({count}x)")
            if len(unique_entities) > 10:
                report_content.append(f"  ... et {len(unique_entities) - 10} autres")
            report_content.append("")
        
        # Sauvegarde le rapport
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(report_content))
            
            print(f"📊 Rapport d'analyse sauvegardé: {output_path}")
            return output_path
        except Exception as e:
            raise Exception(f"Erreur lors de la sauvegarde du rapport: {e}")
    
    def get_document_preview(self, max_chars: int = 1000) -> str:
        """
        Retourne un aperçu du contenu du document
        
        Args:
            max_chars: Nombre maximum de caractères à retourner
            
        Returns:
            str: Aperçu du document
        """
        if not self.extracted_text:
            return "Aucun texte extrait"
        
        preview = self.extracted_text[:max_chars]
        if len(self.extracted_text) > max_chars:
            preview += "\n\n[...texte tronqué...]"
        
        return preview